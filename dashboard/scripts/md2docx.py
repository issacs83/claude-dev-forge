#!/usr/bin/env python3
"""Convert Markdown to DOCX with professional formatting."""
import sys
import re
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def md_to_docx(md_path, docx_path=None):
    if not docx_path:
        docx_path = md_path.rsplit('.', 1)[0] + '.docx'
    
    with open(md_path, 'r', encoding='utf-8') as f:
        content = f.read()
    
    doc = Document()
    
    # Style defaults
    style = doc.styles['Normal']
    style.font.name = 'Malgun Gothic'
    style.font.size = Pt(10)
    style.paragraph_format.space_after = Pt(4)
    
    in_code_block = False
    code_lines = []
    in_table = False
    table_rows = []
    
    for line in content.split('\n'):
        # Code block
        if line.strip().startswith('```'):
            if in_code_block:
                # End code block
                code_text = '\n'.join(code_lines)
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(4)
                p.paragraph_format.space_after = Pt(4)
                run = p.add_run(code_text)
                run.font.name = 'Consolas'
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
                code_lines = []
                in_code_block = False
            else:
                in_code_block = True
            continue
        
        if in_code_block:
            code_lines.append(line)
            continue
        
        # Table
        if '|' in line and line.strip().startswith('|'):
            cells = [c.strip() for c in line.strip().strip('|').split('|')]
            if all(set(c) <= set('-: ') for c in cells):
                continue  # separator row
            table_rows.append(cells)
            in_table = True
            continue
        elif in_table:
            # Flush table
            if table_rows:
                cols = max(len(r) for r in table_rows)
                table = doc.add_table(rows=len(table_rows), cols=cols, style='Light Grid Accent 1')
                for i, row_data in enumerate(table_rows):
                    for j, cell_text in enumerate(row_data):
                        if j < cols:
                            table.rows[i].cells[j].text = cell_text
                            for p in table.rows[i].cells[j].paragraphs:
                                p.style.font.size = Pt(9)
                table_rows = []
            in_table = False
        
        stripped = line.strip()
        
        # Empty line
        if not stripped:
            continue
        
        # Headers
        if stripped.startswith('# '):
            p = doc.add_heading(stripped[2:], level=1)
        elif stripped.startswith('## '):
            p = doc.add_heading(stripped[3:], level=2)
        elif stripped.startswith('### '):
            p = doc.add_heading(stripped[4:], level=3)
        elif stripped.startswith('#### '):
            p = doc.add_heading(stripped[5:], level=4)
        # Bullet list
        elif stripped.startswith('- ') or stripped.startswith('* '):
            doc.add_paragraph(stripped[2:], style='List Bullet')
        elif re.match(r'^\d+\. ', stripped):
            text = re.sub(r'^\d+\. ', '', stripped)
            doc.add_paragraph(text, style='List Number')
        # Bold line
        elif stripped.startswith('**') and stripped.endswith('**'):
            p = doc.add_paragraph()
            run = p.add_run(stripped.strip('*'))
            run.bold = True
        else:
            # Clean markdown formatting
            text = re.sub(r'\*\*(.*?)\*\*', r'\1', stripped)
            text = re.sub(r'\*(.*?)\*', r'\1', text)
            text = re.sub(r'`(.*?)`', r'\1', text)
            doc.add_paragraph(text)
    
    # Flush remaining table
    if table_rows:
        cols = max(len(r) for r in table_rows)
        table = doc.add_table(rows=len(table_rows), cols=cols, style='Light Grid Accent 1')
        for i, row_data in enumerate(table_rows):
            for j, cell_text in enumerate(row_data):
                if j < cols:
                    table.rows[i].cells[j].text = cell_text
    
    doc.save(docx_path)
    print(f'OK: {docx_path}')
    return docx_path

if __name__ == '__main__':
    if len(sys.argv) < 2:
        print('Usage: md2docx.py input.md [output.docx]')
        sys.exit(1)
    out = sys.argv[2] if len(sys.argv) > 2 else None
    md_to_docx(sys.argv[1], out)
